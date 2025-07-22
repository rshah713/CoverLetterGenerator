from docx import Document
from docx.shared import Inches

from datetime import datetime


def replace_doc_text(doc, old_text, new_text):

    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            for run in paragraph.runs:
                run.text = run.text.replace(old_text, new_text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                        if old_text in paragraph.text:
                            for run in paragraph.runs:
                                run.text = run.text.replace(old_text, new_text)
    
    return doc


def handler(company_name, team_name):
    today_str = datetime.now().strftime("%d %B %Y").upper()

    doc = Document("template.docx")
    doc = replace_doc_text(doc, "CompanyName", company_name)
    doc = replace_doc_text(doc, "_date", today_str)
    doc = replace_doc_text(doc, "TeamName", team_name)
    doc.save(f"temp_filled.docx")


if __name__ == "__main__":
    today_str = datetime.now().strftime("%d %B %Y").upper()

    doc = Document("template.docx")
    doc = replace_doc_text(doc, "CompanyName", "Test Company")
    doc = replace_doc_text(doc, "_date", today_str)
    doc = replace_doc_text(doc, "TeamName", "Software Engineering")
    doc.save("Test_Company.docx")
    